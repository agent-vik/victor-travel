#!/usr/bin/env python3
"""Build Word (.docx) travel guides from structured Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from guide_i18n import UI, section_key
from guide_render import (
    BOLD_ONLY_RE,
    DAY_H_RE,
    GEO_H3_RE,
    POI_RE,
    REF_ORIGIN_RE,
    STAY_KEYS,
    TASK_RE,
    parse_day_body,
    parse_trip_header,
    split_days,
    split_sections,
)

FONT_LATIN = "Helvetica Neue"
FONT_CJK = "PingFang SC"
COLOR_PRIMARY = RGBColor(0x2A, 0x9D, 0x8F)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _strip_md_bold(text: str) -> str:
    text = text.strip()
    if text.startswith("**") and text.endswith("**"):
        return text[2:-2].strip()
    return text


def _set_run_font(run, *, size: int = 11, bold: bool = False, color=None) -> None:
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)


def _add_hyperlink(paragraph, text: str, url: str, *, size: int = 11) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2A9D8F")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_CJK)
    rPr.append(rFonts)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _fill_md_runs(paragraph, text: str, *, size: int = 11, bold: bool = False) -> None:
    pos = 0
    for m in MD_LINK_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _set_run_font(run, size=size, bold=bold)
        _add_hyperlink(paragraph, m.group(1), m.group(2), size=size)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, size=size, bold=bold)
    elif pos == 0:
        run = paragraph.add_run(text)
        _set_run_font(run, size=size, bold=bold)


def _add_paragraph(doc: Document, text: str, *, style: str | None = None, size: int = 11, bold: bool = False, space_after: int = 6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    _fill_md_runs(p, text, size=size, bold=bold)
    return p


def _add_bullet(doc: Document, text: str, *, level: int = 0, checked: bool | None = None) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.4)
    prefix = ""
    if checked is True:
        prefix = "☑ "
    elif checked is False:
        prefix = "☐ "
    _fill_md_runs(p, prefix + text, size=10.5)


def _parse_poi_line(line: str) -> tuple[str, str, bool] | None:
    m = POI_RE.match(line)
    if not m:
        return None
    done = m.group(2).lower() == "x"
    raw = m.group(3).strip()
    meta_match = re.search(
        r"(?:（([^（）]+ / [^（）]+ / [^（）]+)）|\(([^()]+ / [^()]+ / [^()]+)\))$",
        raw,
    )
    if meta_match:
        meta_raw = meta_match.group(1) or meta_match.group(2) or ""
        name = raw[: meta_match.start()].strip()
        meta = " / ".join(b.strip() for b in meta_raw.split("/"))
        return name, meta, done
    return raw, "", done


def _add_itinerary(doc: Document, content: str) -> None:
    header, days = split_days(content)
    meta = parse_trip_header(header)

    if meta.get("time"):
        _add_paragraph(doc, _strip_md_bold(meta["time"]), size=12, bold=True, space_after=8)

    for key, val in meta.get("logistics") or []:
        if key in STAY_KEYS and meta.get("stays"):
            _add_paragraph(doc, key, bold=True, space_after=2)
            for stay in meta["stays"]:
                _add_bullet(doc, stay, level=1)
        else:
            sep = ": " if val else ""
            _add_paragraph(doc, f"{key}{sep}{val}", space_after=4)

    doc.add_paragraph()  # spacer

    for title_line, body in days:
        m = DAY_H_RE.match(title_line)
        if not m:
            continue
        badge, rest = m.group(1), m.group(2).strip()
        heading = f"{badge}{rest}"
        _add_paragraph(doc, heading, style="Heading 3", size=12, bold=True, space_after=4)

        parts = parse_day_body(body)
        for t in parts["transport"]:
            p = _add_paragraph(doc, t, size=10, space_after=4)
            for run in p.runs:
                _set_run_font(run, size=10, color=COLOR_MUTED)

        for kind, text in parts["meals"]:
            _add_paragraph(doc, f"{kind}: {text}", size=10.5, space_after=3)

        for spot in parts["spots"]:
            _add_paragraph(doc, spot, bold=True, size=10.5, space_after=4)

        for done, text in parts["todos"]:
            _add_bullet(doc, text, checked=done)

        for para in parts["paragraphs"]:
            _add_paragraph(doc, para, size=10.5, space_after=6)


def _add_bullets_from_text(doc: Document, content: str) -> None:
    lines = [ln.strip() for ln in content.splitlines() if ln.strip() and ln.strip() != "---"]
    if not lines:
        return
    if all(ln.startswith("- ") or TASK_RE.match(ln) for ln in lines):
        for ln in lines:
            tm = TASK_RE.match(ln)
            if tm:
                _add_bullet(doc, tm.group(3).strip(), checked=tm.group(2).lower() == "x")
            elif ln.startswith("- "):
                _add_bullet(doc, ln[2:].strip())
    else:
        _add_paragraph(doc, "\n".join(lines), space_after=6)


def _add_tips(doc: Document, content: str) -> None:
    current: list[str] = []
    sub_title: str | None = None

    def flush() -> None:
        nonlocal current, sub_title
        if sub_title:
            _add_paragraph(doc, sub_title, style="Heading 3", size=11, bold=True, space_after=4)
        if current:
            _add_bullets_from_text(doc, "\n".join(current))
        current = []
        sub_title = None

    for line in content.splitlines():
        if line.startswith("### "):
            flush()
            sub_title = line[4:].strip()
            continue
        if line.strip() in ("", "---"):
            continue
        current.append(line)
    flush()


def _add_pois(doc: Document, content: str, lang: str = "zh") -> None:
    current_group = ""
    current_items: list[tuple[str, str, bool]] = []

    def flush_group() -> None:
        nonlocal current_group, current_items
        if not current_items:
            return
        if current_group:
            _add_paragraph(doc, current_group, bold=True, size=11, space_after=4)
        for name, meta, done in current_items:
            label = name if not meta else f"{name} ({meta})"
            _add_bullet(doc, label, checked=done)
        current_group = ""
        current_items = []

    for line in content.splitlines():
        if not line.strip() or line.strip() == "---":
            continue
        rm = REF_ORIGIN_RE.match(line.strip())
        if rm:
            origin = rm.group(1) or rm.group(2) or rm.group(3) or rm.group(4) or ""
            _add_paragraph(doc, f"{UI[lang]['poi_ref']}{origin}", size=10, space_after=8)
            continue
        gm = GEO_H3_RE.match(line)
        if gm:
            flush_group()
            current_group = gm.group(1).strip()
            continue
        bm = BOLD_ONLY_RE.match(line.strip())
        if bm:
            flush_group()
            current_group = bm.group(1).strip()
            continue
        poi = _parse_poi_line(line)
        if poi:
            current_items.append(poi)
            continue
        if current_items:
            name, meta, done = current_items[-1]
            current_items[-1] = (name + " " + line.strip(), meta, done)

    flush_group()


def _add_generic(doc: Document, content: str) -> None:
    for block in re.split(r"\n\s*\n", content.strip()):
        block = block.strip()
        if not block or block == "---":
            continue
        if block.startswith("- ") or TASK_RE.match(block.splitlines()[0]):
            _add_bullets_from_text(doc, block)
        else:
            _add_paragraph(doc, block, space_after=8)


def build_guide_docx(title: str, body: str, source_url: str = "", *, lang: str = "zh") -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    _add_paragraph(doc, title, style="Title", size=20, bold=True, space_after=6)
    if source_url:
        p = _add_paragraph(doc, source_url, size=9, space_after=14)
        for run in p.runs:
            _set_run_font(run, size=9, color=COLOR_MUTED)

    for sec_title, content in split_sections(body.strip()):
        _add_paragraph(doc, sec_title, style="Heading 2", size=14, bold=True, space_after=8)
        key = section_key(sec_title)
        if key == "itinerary":
            _add_itinerary(doc, content)
        elif key == "checklist":
            _add_bullets_from_text(doc, content) if content.strip().startswith("-") else _add_generic(doc, content)
        elif key == "tips":
            _add_tips(doc, content)
        elif key == "pois":
            _add_pois(doc, content, lang)
        else:
            _add_generic(doc, content)
        doc.add_paragraph()

    return doc


def write_guide_docx(
    title: str, body: str, out_path: Path, *, source_url: str = "", lang: str = "zh"
) -> Path:
    doc = build_guide_docx(title, body, source_url=source_url, lang=lang)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
